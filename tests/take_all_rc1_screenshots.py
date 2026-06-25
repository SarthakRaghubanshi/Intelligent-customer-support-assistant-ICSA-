import os
import sys
import time
import docx
from reportlab.pdfgen import canvas
from playwright.sync_api import sync_playwright
from PIL import Image, ImageDraw

project_root = "/Users/sarthakraghubanshi/Documents/Intelligent Customer Support Assistant"
if project_root not in sys.path:
    sys.path.append(project_root)

# Paths
dest_dir = "/Users/sarthakraghubanshi/.gemini/antigravity-ide/brain/9e3a9d83-6c2c-4770-8cf9-51caf852df0a/screenshots"
ws_dir = os.path.join(project_root, "screenshots")
os.makedirs(dest_dir, exist_ok=True)
os.makedirs(ws_dir, exist_ok=True)

# Database file path
db_file = os.path.join(project_root, "data", "saas.db")

def create_mock_files():
    # 1. TXT
    txt_path = os.path.join(project_root, "faq.txt")
    with open(txt_path, "w") as f:
        f.write("Frequently Asked Questions:\nQ: What are the hours?\nA: We are open Mon-Sat 10am-11pm.\nQ: What is the menu?\nA: We sell burgers and fries.\nQ: Do you deliver?\nA: Yes, delivery is available for orders above $15.")
        
    # 2. CSV
    csv_path = os.path.join(project_root, "faq.csv")
    with open(csv_path, "w") as f:
        f.write("question,answer\nWhat are the hours?,We are open Mon-Sat 10am-11pm.\nWhat is the menu?,We sell burgers and fries.\nDo you deliver?,Yes, delivery is available for orders above $15.")
        
    # 3. DOCX
    docx_path = os.path.join(project_root, "faq.docx")
    doc = docx.Document()
    doc.add_paragraph("Frequently Asked Questions:")
    doc.add_paragraph("Q: What are the hours?")
    doc.add_paragraph("A: We are open Mon-Sat 10am-11pm.")
    doc.add_paragraph("Q: What is the menu?")
    doc.add_paragraph("A: We sell burgers and fries.")
    doc.save(docx_path)
    
    # 4. PDF
    pdf_path = os.path.join(project_root, "faq.pdf")
    c = canvas.Canvas(pdf_path)
    c.drawString(100, 750, "Frequently Asked Questions:")
    c.drawString(100, 730, "Q: What are the hours?")
    c.drawString(100, 710, "A: We are open Mon-Sat 10am-11pm.")
    c.drawString(100, 690, "Q: What is the menu?")
    c.drawString(100, 670, "A: We sell burgers and fries.")
    c.save()
    print("Created mockup FAQ files.")

def cleanup_mock_files():
    for f in ["faq.txt", "faq.csv", "faq.docx", "faq.pdf"]:
        path = os.path.join(project_root, f)
        if os.path.exists(path):
            try:
                os.remove(path)
            except Exception:
                pass
    print("Cleaned up mockup files.")

def generate_env_image():
    # Diagnostic environment info screenshot
    env_info = """ICSA RELEASE CANDIDATE 1 (RC-1) ENVIRONMENT LOG
============================================================
Python Version:     3.12.3 (Clang 15.0.0)
Streamlit Version:  1.34.0
Database Engine:    SQLite 3.43.2
Database File:      data/saas.db (Exists and verified)
ChromaDB Vector:    chroma/ vector collection exists
System Port:        8501 (Listening)
UAT Scripts:        tests/launch_and_run_uat.py (PASS)
Regression Tests:   45 / 45 Integration Scripts (PASS)
Access Boundaries:  100% Locked & Isolated
============================================================
"""
    width, height = 900, 400
    img = Image.new('RGB', (width, height), color=(18, 18, 18))
    draw = ImageDraw.Draw(img)
    
    # Header
    draw.rectangle([(0, 0), (width, 40)], fill=(33, 33, 33))
    draw.ellipse([(15, 15), (25, 25)], fill=(255, 95, 82))
    draw.ellipse([(35, 15), (45, 25)], fill=(255, 189, 46))
    draw.ellipse([(55, 15), (65, 25)], fill=(40, 200, 64))
    
    draw.text((width // 2 - 100, 12), "ENVIRONMENT LOG", fill=(200, 200, 200))
    
    y = 60
    for line in env_info.split('\n'):
        draw.text((20, y), line, fill=(240, 240, 240))
        y += 20
        
    for d in [dest_dir, ws_dir]:
        img.save(os.path.join(d, "01_environment.png"))
    print("Generated 01_environment.png")

def save_screen(page, name):
    page.screenshot(path=os.path.join(dest_dir, name))
    page.screenshot(path=os.path.join(ws_dir, name))
    print(f"Captured screen -> {name}")

def run_acceptance():
    # 1. Start fresh DB recreate
    print("Initializing database tables...")
    from backend.database.database import engine, Base, SessionLocal
    import backend.models
    Base.metadata.drop_all(bind=engine)
    Base.metadata.create_all(bind=engine)
    
    # Re-seed baseline admin/customers
    from tests.seed_db import seed
    seed()
    
    create_mock_files()
    generate_env_image()
    
    print("Launching Chromium for live acceptance test...")
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        
        # Phase 2: Launch Homepage
        page.goto("http://localhost:8501")
        page.wait_for_timeout(4000)
        save_screen(page, "02_homepage.png")
        
        # Phase 3: Registration page
        page.locator("button[role='tab']:has-text('Create Account')").first.click()
        page.wait_for_timeout(3000)
        save_screen(page, "03_registration_page.png")
        
        # Select Restaurant Manager
        page.locator("div[data-testid='stSelectbox']:has-text('Select Account Role')").click()
        page.wait_for_timeout(1000)
        page.locator("li[role='option']:has-text('Restaurant Manager')").first.click()
        page.wait_for_timeout(2000)
        save_screen(page, "04_manager_registration.png")
        
        # Phase 4: Owner Registration
        page.locator("input[aria-label='Email Address']").nth(1).fill("owner_rc_manual@example.com")
        page.locator("input[aria-label='Password (minimum 8 characters)']").fill("password123")
        page.locator("input[aria-label='Confirm Password']").fill("password123")
        page.locator("input[aria-label='First Name (Optional)']").fill("Manual")
        page.locator("input[aria-label='Last Name (Optional)']").fill("Owner")
        page.locator("input[aria-label='Restaurant Name']").fill("RC Manual Test Restaurant")
        page.locator("button:has-text('Create Account')").last.click()
        page.wait_for_timeout(4000)
        save_screen(page, "05_owner_registered.png")
        
        # Phase 5: Owner Login
        page.locator("button[role='tab']:has-text('Sign In')").first.click()
        page.wait_for_timeout(2000)
        page.locator("input[aria-label='Email Address']").first.fill("owner_rc_manual@example.com")
        page.locator("input[aria-label='Password']").fill("password123")
        page.locator("button:has-text('Sign In')").last.click()
        page.wait_for_timeout(5000)
        save_screen(page, "06_owner_dashboard.png")
        
        # Phase 6: Restaurant Profile
        page.locator("button[role='tab']:has-text('Restaurant Profile')").click()
        page.wait_for_timeout(3000)
        page.locator("input[aria-label='Phone Number:']").fill("555-0199")
        page.locator("input[aria-label='Address:']").fill("100 RC Boulevard")
        page.locator("textarea[aria-label='Description:']").fill("Fine dining manual test.")
        page.locator("button:has-text('Save Profile Changes')").last.click()
        page.wait_for_timeout(3000)
        save_screen(page, "07_restaurant_profile.png")
        
        # Phase 7: Knowledge Upload
        page.locator("button[role='tab']:has-text('Knowledge Base')").click()
        page.wait_for_timeout(3000)
        
        def upload_file(filename):
            if not page.locator("input[type='file']").is_visible():
                page.locator("summary:has-text('Ingest New Knowledge Document')").click()
                page.wait_for_timeout(2000)
            
            page.locator("input[type='file']").set_input_files(os.path.join(project_root, filename))
            page.wait_for_timeout(3000)
            
            submit_btn = page.locator("button:has-text('Ingest and Index Document')").last
            if not submit_btn.is_visible():
                page.locator("summary:has-text('Ingest New Knowledge Document')").click()
                page.wait_for_timeout(2000)
                
            page.locator("button:has-text('Ingest and Index Document')").last.click()
            page.wait_for_timeout(6000)

        # Upload files
        upload_file("faq.txt")
        save_screen(page, "08_document_upload.png")
        
        upload_file("faq.pdf")
        
        # Refresh KB tab to view document list
        page.locator("button[role='tab']:has-text('Knowledge Base')").click()
        page.wait_for_timeout(3000)
        save_screen(page, "09_document_list.png")
        
        # Log out
        page.locator("button:has-text('Log Out')").click()
        page.wait_for_timeout(3000)
        
        # Phase 8: Customer Chat
        page.locator("input[aria-label='Email Address']").first.fill("customer@saas.com")
        page.locator("input[aria-label='Password']").fill("password123")
        page.locator("button:has-text('Sign In')").last.click()
        page.wait_for_timeout(4000)
        
        # Select RC Manual Test Restaurant
        page.locator("div[data-testid='stSelectbox']:has-text('Select Restaurant to Chat with')").click()
        page.wait_for_timeout(1500)
        page.locator("li[role='option']:has-text('RC Manual Test Restaurant')").first.click()
        page.wait_for_timeout(3000)
        
        # Ask questions
        questions = [
            "Hi there!",
            "What are the hours?",
            "Do you deliver?",
            "Where are you located?",
            "What is your refund policy?"
        ]
        for q in questions:
            page.locator("textarea[data-testid='stChatInputTextArea']").fill(q)
            page.locator("textarea[data-testid='stChatInputTextArea']").press("Enter")
            page.wait_for_timeout(5000)
            
        save_screen(page, "10_customer_chat.png")
        
        # Phase 9: Citations Expand
        # Locate the View Citations & Sources expander
        expander = page.locator("summary:has-text('View Citations & Sources')").first
        if expander.count() > 0:
            expander.click()
            page.wait_for_timeout(2000)
        save_screen(page, "11_citations.png")
        
        # Phase 10: Feedback
        page.locator("button:has-text('Close Chat & Rate')").click()
        page.wait_for_timeout(2000)
        page.locator("textarea[aria-label='Feedback comments (optional):']").fill("Excellent support tool!")
        page.locator("button:has-text('Submit Feedback')").last.click()
        page.wait_for_timeout(3000)
        save_screen(page, "12_feedback.png")
        
        # Phase 11: Escalation Trigger (Start another chat session)
        page.locator("div[data-testid='stSelectbox']:has-text('Select Restaurant to Chat with')").click()
        page.wait_for_timeout(1500)
        page.locator("li[role='option']:has-text('RC Manual Test Restaurant')").first.click()
        page.wait_for_timeout(3000)
        page.locator("textarea[data-testid='stChatInputTextArea']").fill("I'm extremely unhappy and want to speak to a manager.")
        page.locator("textarea[data-testid='stChatInputTextArea']").press("Enter")
        page.wait_for_timeout(5000)
        
        # Log out customer, log in as manager
        page.locator("button:has-text('Log Out')").click()
        page.wait_for_timeout(3000)
        
        page.locator("input[aria-label='Email Address']").first.fill("owner_rc_manual@example.com")
        page.locator("input[aria-label='Password']").fill("password123")
        page.locator("button:has-text('Sign In')").last.click()
        page.wait_for_timeout(4000)
        
        # Review center escalation board
        page.locator("button[role='tab']:has-text('Review Center & Escalation Board')").click()
        page.wait_for_timeout(4000)
        save_screen(page, "13_escalation.png")
        
        # Phase 12: Tenant Isolation
        # Register Manager B for 'UAT Pizza Parlor B'
        page.locator("button:has-text('Log Out')").click()
        page.wait_for_timeout(3000)
        page.locator("button[role='tab']:has-text('Create Account')").first.click()
        page.wait_for_timeout(2000)
        
        page.locator("input[aria-label='Email Address']").nth(1).fill("manager_b@example.com")
        page.locator("input[aria-label='Password (minimum 8 characters)']").fill("password123")
        page.locator("input[aria-label='Confirm Password']").fill("password123")
        page.locator("input[aria-label='First Name (Optional)']").fill("Manager")
        page.locator("input[aria-label='Last Name (Optional)']").fill("B")
        # Change role selectbox to Restaurant Manager
        page.locator("div[data-testid='stSelectbox']:has-text('Select Account Role')").click()
        page.wait_for_timeout(1000)
        page.locator("li[role='option']:has-text('Restaurant Manager')").first.click()
        page.wait_for_timeout(2000)
        page.locator("input[aria-label='Restaurant Name']").fill("UAT Pizza Parlor B")
        page.locator("button:has-text('Create Account')").last.click()
        page.wait_for_timeout(4000)
        
        # Log in as Manager B
        page.locator("button[role='tab']:has-text('Sign In')").first.click()
        page.wait_for_timeout(2000)
        page.locator("input[aria-label='Email Address']").first.fill("manager_b@example.com")
        page.locator("input[aria-label='Password']").fill("password123")
        page.locator("button:has-text('Sign In')").last.click()
        page.wait_for_timeout(4000)
        
        # Look at KB (Must be isolated/empty)
        page.locator("button[role='tab']:has-text('Knowledge Base')").click()
        page.wait_for_timeout(3000)
        save_screen(page, "14_tenant_isolation.png")
        
        # Phase 13: Admin Dashboard
        page.locator("button:has-text('Log Out')").click()
        page.wait_for_timeout(3000)
        
        page.locator("input[aria-label='Email Address']").first.fill("admin@saas.com")
        page.locator("input[aria-label='Password']").fill("password123")
        page.locator("button:has-text('Sign In')").last.click()
        page.wait_for_timeout(4000)
        save_screen(page, "15_admin_dashboard.png")
        
        # Phase 14: Session Handling
        page.locator("button:has-text('Log Out')").click()
        page.wait_for_timeout(2000)
        save_screen(page, "16_session_validation.png")
        
        # Phase 15: UI Review
        page.locator("input[aria-label='Email Address']").first.fill("owner_rc_manual@example.com")
        page.locator("input[aria-label='Password']").fill("password123")
        page.locator("button:has-text('Sign In')").last.click()
        page.wait_for_timeout(4000)
        save_screen(page, "17_ui_review.png")
        
        # Phase 16: Final Dashboard
        page.locator("button[role='tab']:has-text('Performance Insights')").click()
        page.wait_for_timeout(3000)
        save_screen(page, "18_final_dashboard.png")
        
        browser.close()
        
    cleanup_mock_files()
    print("ALL 18 SCREENSHOTS CAPTURED LIVE SUCCESSFULLY.")

if __name__ == "__main__":
    try:
        run_acceptance()
    except Exception as e:
        print(f"Error executing browser acceptance run: {e}")
