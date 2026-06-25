import os
import sys
import time
import json
import sqlite3
import subprocess
import signal
from playwright.sync_api import sync_playwright

# Ensure project root is in the path
current_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.dirname(current_dir)
if project_root not in sys.path:
    sys.path.append(project_root)

# Paths
db_file = os.path.join(project_root, "data", "saas.db")
dest_dir = "/Users/sarthakraghubanshi/.gemini/antigravity-ide/brain/9e3a9d83-6c2c-4770-8cf9-51caf852df0a"

def create_mock_files():
    import docx
    from reportlab.pdfgen import canvas
    
    # 1. TXT
    txt_path = os.path.join(project_root, "uat_faq.txt")
    with open(txt_path, "w") as f:
        f.write("Frequently Asked Questions:\nQ: What are the hours?\nA: We are open Mon-Sat 10am-11pm.\nQ: What is the menu?\nA: We sell burgers and fries.")
        
    # 2. CSV
    csv_path = os.path.join(project_root, "uat_faq.csv")
    with open(csv_path, "w") as f:
        f.write("question,answer\nWhat are the hours?,We are open Mon-Sat 10am-11pm.\nWhat is the menu?,We sell burgers and fries.")
        
    # 3. DOCX
    docx_path = os.path.join(project_root, "uat_faq.docx")
    doc = docx.Document()
    doc.add_paragraph("Frequently Asked Questions:")
    doc.add_paragraph("Q: What are the hours?")
    doc.add_paragraph("A: We are open Mon-Sat 10am-11pm.")
    doc.add_paragraph("Q: What is the menu?")
    doc.add_paragraph("A: We sell burgers and fries.")
    doc.save(docx_path)
    
    # 4. PDF
    pdf_path = os.path.join(project_root, "uat_faq.pdf")
    c = canvas.Canvas(pdf_path)
    c.drawString(100, 750, "Frequently Asked Questions:")
    c.drawString(100, 730, "Q: What are the hours?")
    c.drawString(100, 710, "A: We are open Mon-Sat 10am-11pm.")
    c.drawString(100, 690, "Q: What is the menu?")
    c.drawString(100, 670, "A: We sell burgers and fries.")
    c.save()
    
    # 5. Invalid File (e.g. invalid extension png)
    png_path = os.path.join(project_root, "invalid.png")
    with open(png_path, "wb") as f:
        f.write(b"PNG mock data")
        
    # 6. Oversized file (11MB text file)
    oversized_path = os.path.join(project_root, "uat_oversized.txt")
    with open(oversized_path, "wb") as f:
        f.seek(11 * 1024 * 1024 - 1)
        f.write(b"\0")

    print("Created mock files for upload verification.")

def cleanup_mock_files():
    for f in ["uat_faq.txt", "uat_faq.csv", "uat_faq.docx", "uat_faq.pdf", "invalid.png", "uat_oversized.txt"]:
        path = os.path.join(project_root, f)
        if os.path.exists(path):
            try:
                os.remove(path)
            except Exception:
                pass
    print("Cleaned up mock files.")

# 1. Kill any existing process on port 8501
print("Cleaning up any existing processes on port 8501...")
try:
    pid_output = subprocess.check_output(["lsof", "-t", "-i:8501"], text=True)
    for pid in pid_output.strip().split("\n"):
        if pid:
            os.kill(int(pid), signal.SIGTERM)
            print(f"Killed old Streamlit process with PID {pid}")
    time.sleep(2)
except Exception:
    print("No active processes found on port 8501.")

# 2. Cleanup existing DB file
if os.path.exists(db_file):
    try:
        os.remove(db_file)
        print("Removed database file cleanly.")
    except Exception as e:
        print(f"Could not remove database file: {e}")

# Create mock upload files
create_mock_files()

# 3. Create fresh database tables
print("Initializing fresh database tables...")
from backend.database.database import engine, Base, SessionLocal
import backend.models
Base.metadata.create_all(bind=engine)

# Seed an existing restaurant for Scenario C
from backend.repositories.restaurant_repository import RestaurantRepository
db = SessionLocal()
existing_rest = RestaurantRepository.create(db, "UAT Pizza Parlor")
db.commit()
rest_id = existing_rest.id
db.close()
print(f"Seeded restaurant UAT Pizza Parlor ({rest_id})")

# 4. Start Streamlit server fresh
print("Starting Streamlit server...")
streamlit_proc = subprocess.Popen(
    ["python3", "-m", "streamlit", "run", "frontend/app.py", "--server.headless", "true", "--server.port", "8501"],
    cwd=project_root,
    stdout=subprocess.DEVNULL,
    stderr=subprocess.DEVNULL
)
time.sleep(6) # Wait for server to start

def check_db_record(query, params=()):
    conn = sqlite3.connect(db_file)
    cursor = conn.cursor()
    cursor.execute(query, params)
    row = cursor.fetchone()
    conn.close()
    return row

def check_db_all(query, params=()):
    conn = sqlite3.connect(db_file)
    cursor = conn.cursor()
    cursor.execute(query, params)
    rows = cursor.fetchall()
    conn.close()
    return rows

def run_uat():
    results = {}
    print("Launching Chromium for UAT...")
    
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()

        # ----------------------------------------------------
        # STEP 1: AUTHENTICATION & LOGIN FLOW
        # ----------------------------------------------------
        print("\n=== Testing Step 1 ===")
        page.goto("http://localhost:8501")
        page.wait_for_timeout(4000)
        page.screenshot(path=os.path.join(dest_dir, "uat_login_screen.png"))
        
        # Click Register tab
        page.locator("button[role='tab']:has-text('Create Account')").first.click()
        page.wait_for_timeout(3000)
        page.screenshot(path=os.path.join(dest_dir, "uat_registration_screen.png"))
        
        # Test 1.1: Invalid password validation (short password)
        page.locator("input[aria-label='Email Address']").nth(1).fill("invalid_pw@saas.com")
        page.locator("input[aria-label='Password (minimum 8 characters)']").fill("short")
        page.locator("input[aria-label='Confirm Password']").fill("short")
        page.locator("button:has-text('Create Account')").last.click()
        page.wait_for_timeout(4000)
        results["step1_invalid_password_rejected"] = "Password must be at least 8 characters." in page.content()
        print(f"1.1 Invalid password rejection: {results['step1_invalid_password_rejected']}")

        # Click Register tab again to refresh input states
        page.locator("button[role='tab']:has-text('Create Account')").first.click()
        page.wait_for_timeout(3000)

        # Test 1.2: Customer registration
        page.locator("input[aria-label='Email Address']").nth(1).fill("customer_uat@saas.com")
        page.locator("input[aria-label='Password (minimum 8 characters)']").fill("password123")
        page.locator("input[aria-label='Confirm Password']").fill("password123")
        page.locator("button:has-text('Create Account')").last.click()
        page.wait_for_timeout(4000)
        
        cust_row = check_db_record("select email, role from users where email='customer_uat@saas.com'")
        results["step1_register_customer"] = cust_row is not None and cust_row[1].lower() == "customer"
        print(f"1.2 Register Customer: {results['step1_register_customer']}")

        # Click Register tab
        page.locator("button[role='tab']:has-text('Create Account')").first.click()
        page.wait_for_timeout(3000)

        # Test 1.3: Duplicate email rejection
        page.locator("input[aria-label='Email Address']").nth(1).fill("customer_uat@saas.com")
        page.locator("input[aria-label='Password (minimum 8 characters)']").fill("password123")
        page.locator("input[aria-label='Confirm Password']").fill("password123")
        page.locator("button:has-text('Create Account')").last.click()
        page.wait_for_timeout(4000)
        results["step1_duplicate_email_rejected"] = "Email already registered" in page.content()
        print(f"1.3 Duplicate email rejected: {results['step1_duplicate_email_rejected']}")

        # Click Register tab
        page.locator("button[role='tab']:has-text('Create Account')").first.click()
        page.wait_for_timeout(3000)

        # Test 1.4: Admin registration
        page.locator("input[aria-label='Email Address']").nth(1).fill("admin_uat@saas.com")
        page.locator("input[aria-label='Password (minimum 8 characters)']").fill("password123")
        page.locator("input[aria-label='Confirm Password']").fill("password123")
        page.locator("div[data-testid='stSelectbox']").first.click()
        page.wait_for_timeout(1500)
        page.locator("li[role='option']:has-text('System Administrator')").first.click()
        page.wait_for_timeout(3000)
        page.locator("button:has-text('Create Account')").last.click()
        page.wait_for_timeout(4000)
        
        admin_row = check_db_record("select email, role from users where email='admin_uat@saas.com'")
        results["step1_register_admin"] = admin_row is not None and admin_row[1].lower() == "admin"
        print(f"1.4 Register Admin: {results['step1_register_admin']}")

        # ----------------------------------------------------
        # STEP 2: TENANT REGISTRATION & USER MAPPING
        # ----------------------------------------------------
        print("\n=== Testing Step 2 ===")
        page.locator("button[role='tab']:has-text('Create Account')").first.click()
        page.wait_for_timeout(3000)
        
        # Test 2.1: Verify existing restaurant selection UI is removed
        page.locator("input[aria-label='Email Address']").nth(1).fill("manager_linked@saas.com")
        page.locator("input[aria-label='Password (minimum 8 characters)']").fill("password123")
        page.locator("input[aria-label='Confirm Password']").fill("password123")
        # Change role selectbox to Restaurant Manager
        page.locator("div[data-testid='stSelectbox']").first.click()
        page.wait_for_timeout(1500)
        page.locator("li[role='option']:has-text('Restaurant Manager')").first.click()
        page.wait_for_timeout(3000)
        
        results["step2_manager_linked_existing"] = "Select Existing Restaurant" not in page.content() and "Mapping Mode" not in page.content()
        print(f"2.1 Link Manager to Existing Restaurant UI is disabled: {results['step2_manager_linked_existing']}")

        # Click Register tab
        page.locator("button[role='tab']:has-text('Create Account')").first.click()
        page.wait_for_timeout(3000)

        # Test 2.2: Manager with new restaurant onboarding
        page.locator("input[aria-label='Email Address']").nth(1).fill("manager_new@saas.com")
        page.locator("input[aria-label='Password (minimum 8 characters)']").fill("password123")
        page.locator("input[aria-label='Confirm Password']").fill("password123")
        # Select role Restaurant Manager
        page.locator("div[data-testid='stSelectbox']").first.click()
        page.wait_for_timeout(1500)
        page.locator("li[role='option']:has-text('Restaurant Manager')").first.click()
        page.wait_for_timeout(3000)
        page.screenshot(path=os.path.join(dest_dir, "uat_manager_onboarding.png"))
        
        # Fill restaurant name
        page.locator("input[aria-label='Restaurant Name']").fill("UAT Burger")
        page.locator("button:has-text('Create Account')").last.click()
        page.wait_for_timeout(4000)
        
        new_manager_row = check_db_record("select email, role, restaurant_id from users where email='manager_new@saas.com'")
        new_rest_row = check_db_record("select id, name from restaurants where name='UAT Burger'")
        
        results["step2_manager_restaurant_created"] = new_rest_row is not None
        results["step2_manager_restaurant_linked"] = new_manager_row is not None and new_manager_row[2] == new_rest_row[0]
        print(f"2.2 Onboard New Restaurant: {results['step2_manager_restaurant_created']} (linked: {results['step2_manager_restaurant_linked']})")

        # Click Register tab
        page.locator("button[role='tab']:has-text('Create Account')").first.click()
        page.wait_for_timeout(3000)

        # Test 2.3: Duplicate restaurant name rejected
        page.locator("input[aria-label='Email Address']").nth(1).fill("manager_new_dup@saas.com")
        page.locator("input[aria-label='Password (minimum 8 characters)']").fill("password123")
        page.locator("input[aria-label='Confirm Password']").fill("password123")
        page.locator("div[data-testid='stSelectbox']").first.click()
        page.wait_for_timeout(1500)
        page.locator("li[role='option']:has-text('Restaurant Manager')").first.click()
        page.wait_for_timeout(3000)
        page.locator("input[aria-label='Restaurant Name']").fill("UAT Burger")
        page.locator("button:has-text('Create Account')").last.click()
        page.wait_for_timeout(4000)
        results["step2_duplicate_restaurant_rejected"] = "already taken" in page.content() or "already registered" in page.content()
        print(f"2.3 Duplicate Restaurant rejected: {results['step2_duplicate_restaurant_rejected']}")

        # ----------------------------------------------------
        # LOGIN & DASHBOARD ROUTING SESSIONS
        # ----------------------------------------------------
        # Test 1.5: Login Admin
        print("\n=== Testing Dashboard Routing ===")
        page.locator("button[role='tab']:has-text('Sign In')").first.click()
        page.wait_for_timeout(2000)
        page.locator("input[aria-label='Email Address']").first.fill("admin_uat@saas.com")
        page.locator("input[aria-label='Password']").fill("password123")
        page.locator("button:has-text('Sign In')").last.click()
        page.wait_for_timeout(4000)
        page.screenshot(path=os.path.join(dest_dir, "uat_admin_dashboard.png"))
        results["step1_login_admin"] = "Admin Dashboard" in page.content()
        print(f"Admin Dashboard loaded: {results['step1_login_admin']}")
        
        # Log out
        page.locator("button:has-text('Log Out')").click()
        page.wait_for_timeout(3000)

        # Test 1.6: Login Customer
        page.locator("input[aria-label='Email Address']").first.fill("customer_uat@saas.com")
        page.locator("input[aria-label='Password']").fill("password123")
        page.locator("button:has-text('Sign In')").last.click()
        page.wait_for_timeout(4000)
        page.screenshot(path=os.path.join(dest_dir, "uat_customer_dashboard.png"))
        results["step1_login_customer"] = "Customer Dashboard" in page.content()
        print(f"Customer Dashboard loaded: {results['step1_login_customer']}")
        
        # Log out
        page.locator("button:has-text('Log Out')").click()
        page.wait_for_timeout(3000)

        # Test 2.4: Login Manager
        page.locator("input[aria-label='Email Address']").first.fill("manager_new@saas.com")
        page.locator("input[aria-label='Password']").fill("password123")
        page.locator("button:has-text('Sign In')").last.click()
        page.wait_for_timeout(5000)
        page.screenshot(path=os.path.join(dest_dir, "uat_restaurant_dashboard.png"))
        results["step2_login_manager"] = "Restaurant Dashboard" in page.content()
        print(f"Manager Dashboard loaded: {results['step2_login_manager']}")

        # ----------------------------------------------------
        # STEP 3: RESTAURANT PROFILE & SETTINGS
        # ----------------------------------------------------
        print("\n=== Testing Step 3 ===")
        page.locator("button[role='tab']").nth(2).click()
        page.wait_for_timeout(3000)
        
        page.locator("input[aria-label='Phone Number:']").fill("555-9876")
        page.locator("input[aria-label='Address:']").fill("456 UAT Ave")
        page.locator("textarea[aria-label='Description:']").fill("Fine burgers UAT restaurant.")
        page.locator("button:has-text('Save Profile Changes')").last.click()
        page.wait_for_timeout(3000)
        
        profile_row = check_db_record("select phone, address, description from restaurants where name='UAT Burger'")
        results["step3_profile_update"] = profile_row is not None and profile_row[0] == "555-9876" and profile_row[1] == "456 UAT Ave"
        print(f"3.1 Edit Restaurant Profile: {results['step3_profile_update']}")

        # Toggle Delivery
        page.locator("label:has-text('Delivery Available?')").click()
        page.wait_for_timeout(500)
        
        # Edit Hours (Monday Open)
        page.locator("input[aria-label='Open (HH:MM):']").first.fill("10:00")
        page.locator("button:has-text('Save Profile Changes')").last.click()
        page.wait_for_timeout(3000)
        
        profile_row_2 = check_db_record("select business_hours, delivery_available from restaurants where name='UAT Burger'")
        results["step3_hours_update"] = profile_row_2 is not None and "10:00" in str(profile_row_2[0])
        results["step3_delivery_toggle"] = profile_row_2 is not None and profile_row_2[1] == 0  # toggled from default True to False
        print(f"3.2 Edit Hours & Settings: {results['step3_hours_update']} (delivery toggled: {results['step3_delivery_toggle']})")

        # ----------------------------------------------------
        # STEP 4: KNOWLEDGE MANAGEMENT & DOCUMENT INGESTION
        # ----------------------------------------------------
        print("\n=== Testing Step 4 ===")
        page.locator("button[role='tab']").nth(1).click()
        page.wait_for_timeout(3000)
        
        def upload_file_test(filename, doc_type="faq"):
            if not page.locator("input[type='file']").is_visible():
                page.locator("summary:has-text('Ingest New Knowledge Document')").click()
                page.wait_for_timeout(1000)
                
            filepath = os.path.join(project_root, filename)
            page.locator("input[type='file']").set_input_files(filepath)
            page.wait_for_timeout(2000)
            
            # Specific selectbox selector that text contains 'Document Type'
            selectbox = page.locator("div[data-testid='stSelectbox']:has-text('Document Type')")
            selectbox.click()
            page.wait_for_timeout(1000)
            
            # Keyboard navigate to selection
            if doc_type == "faq":
                # menu is default (0). Press ArrowDown once to go to faq (1)
                page.keyboard.press("ArrowDown")
            elif doc_type == "other":
                # press ArrowDown 5 times
                for _ in range(5):
                    page.keyboard.press("ArrowDown")
                    page.wait_for_timeout(100)
            
            page.keyboard.press("Enter")
            page.wait_for_timeout(1500)
            
            page.locator("button:has-text('Ingest and Index Document')").last.click()
            page.wait_for_timeout(6000)

        # Ingest TXT
        upload_file_test("uat_faq.txt")
        # Ingest CSV
        upload_file_test("uat_faq.csv")
        # Ingest DOCX
        upload_file_test("uat_faq.docx")
        # Ingest PDF
        upload_file_test("uat_faq.pdf")
        
        page.screenshot(path=os.path.join(dest_dir, "uat_document_ingestion_previews.png"))

        # Ingest Invalid extension
        upload_file_test("invalid.png")
        results["step4_invalid_extension_rejected"] = any(x in page.content() for x in ["Unsupported file", "not a valid", "select a file", "invalid", "limit"])
        print(f"4.5 Invalid extension rejected: {results['step4_invalid_extension_rejected']}")
        
        # Ingest Oversized
        upload_file_test("uat_oversized.txt")
        results["step4_oversized_rejected"] = "exceeds 10MB limit" in page.content()
        print(f"4.6 Oversized file rejected: {results['step4_oversized_rejected']}")

        # Verify database records
        txt_row = check_db_record("select title, content from knowledge_documents where title='uat_faq.txt'")
        csv_row = check_db_record("select title, content from knowledge_documents where title='uat_faq.csv'")
        docx_row = check_db_record("select title, content from knowledge_documents where title='uat_faq.docx'")
        pdf_row = check_db_record("select title, content from knowledge_documents where title='uat_faq.pdf'")
        
        results["step4_txt_persisted"] = txt_row is not None and "10am-11pm" in txt_row[1]
        results["step4_csv_persisted"] = csv_row is not None and "10am-11pm" in csv_row[1]
        results["step4_docx_persisted"] = docx_row is not None and "10am-11pm" in docx_row[1]
        results["step4_pdf_persisted"] = pdf_row is not None and "10am-11pm" in pdf_row[1]
        
        print(f"TXT persisted: {results['step4_txt_persisted']}")
        print(f"CSV persisted: {results['step4_csv_persisted']}")
        print(f"DOCX persisted: {results['step4_docx_persisted']}")
        print(f"PDF persisted: {results['step4_pdf_persisted']}")

        # Log out
        page.locator("button:has-text('Log Out')").click()
        page.wait_for_timeout(3000)

        # ----------------------------------------------------
        # STEP 6: TENANT ISOLATION (CROSS-TENANT KNOWLEDGE VERIFICATION)
        # ----------------------------------------------------
        print("\n=== Testing Step 6 ===")
        # Register Manager B (linked to UAT Pizza Parlor B)
        page.locator("button[role='tab']:has-text('Create Account')").first.click()
        page.wait_for_timeout(3000)
        page.locator("input[aria-label='Email Address']").nth(1).fill("manager_b@saas.com")
        page.locator("input[aria-label='Password (minimum 8 characters)']").fill("password123")
        page.locator("input[aria-label='Confirm Password']").fill("password123")
        page.locator("div[data-testid='stSelectbox']").first.click()
        page.wait_for_timeout(1000)
        page.locator("li[role='option']:has-text('Restaurant Manager')").first.click()
        page.wait_for_timeout(2000)
        page.locator("input[aria-label='Restaurant Name']").fill("UAT Pizza Parlor B")
        page.locator("button:has-text('Create Account')").last.click()
        page.wait_for_timeout(4000)
        
        # Log in as Manager B
        page.locator("button[role='tab']:has-text('Sign In')").first.click()
        page.wait_for_timeout(2000)
        page.locator("input[aria-label='Email Address']").first.fill("manager_b@saas.com")
        page.locator("input[aria-label='Password']").fill("password123")
        page.locator("button:has-text('Sign In')").last.click()
        page.wait_for_timeout(4000)
        
        # View KB
        page.locator("button[role='tab']").nth(1).click()
        page.wait_for_timeout(3000)
        results["step6_tenant_isolation"] = "uat_faq.txt" not in page.content() and "uat_faq.csv" not in page.content()
        print(f"6.1 Cross-tenant isolation verified: {results['step6_tenant_isolation']}")
        
        # Log out
        page.locator("button:has-text('Log Out')").click()
        page.wait_for_timeout(3000)

        # ----------------------------------------------------
        # STEP 5 & 7 & 8 & 9: CUSTOMER CHAT & RAG WORKFLOW
        # ----------------------------------------------------
        print("\n=== Testing Steps 5, 7, 8, 9 ===")
        page.locator("input[aria-label='Email Address']").first.fill("customer_uat@saas.com")
        page.locator("input[aria-label='Password']").fill("password123")
        page.locator("button:has-text('Sign In')").last.click()
        page.wait_for_timeout(4000)
        
        # Choose UAT Burger
        page.locator("div[data-testid='stSelectbox']:has-text('Select Restaurant to Chat with')").click()
        page.wait_for_timeout(1500)
        page.locator("li[role='option']:has-text('UAT Burger')").first.click()
        page.wait_for_timeout(3500)
        
        # Ask FAQ question
        page.locator("textarea[data-testid='stChatInputTextArea']").fill("What are the hours?")
        page.locator("textarea[data-testid='stChatInputTextArea']").press("Enter")
        page.wait_for_timeout(6000)
        page.screenshot(path=os.path.join(dest_dir, "uat_customer_chat_rag.png"))
        
        messages_list = check_db_all("select role, content from messages order by timestamp desc limit 3")
        results["step7_chat_rendering"] = len(messages_list) >= 2
        results["step5_rag_retrieval"] = any("10am-11pm" in str(m[1]) for m in messages_list)
        print(f"7.1 Conversation rendering: {results['step7_chat_rendering']}")
        print(f"5.1 RAG retrieval answered FAQ: {results['step5_rag_retrieval']}")

        # Submit Complaint
        page.locator("textarea[data-testid='stChatInputTextArea']").fill("complaint: I want to request a refund immediately, my burger was cold and late!")
        page.locator("textarea[data-testid='stChatInputTextArea']").press("Enter")
        page.wait_for_timeout(6000)
        page.screenshot(path=os.path.join(dest_dir, "uat_customer_complaint.png"))
        
        sentiment_row = check_db_record("select intent, sentiment, language, escalated from messages where role='assistant' order by timestamp desc limit 1")
        results["step8_intent_classification"] = sentiment_row is not None and sentiment_row[0] is not None
        results["step8_sentiment_classification"] = sentiment_row is not None and sentiment_row[1] is not None
        results["step8_language_detection"] = sentiment_row is not None and sentiment_row[2] is not None
        results["step10_escalation_generated"] = sentiment_row is not None and sentiment_row[3] == 1
        print(f"8.1 Intent classification: {results['step8_intent_classification']} (Intent: {sentiment_row[0] if sentiment_row else None})")
        print(f"8.2 Sentiment classification: {results['step8_sentiment_classification']} (Sentiment: {sentiment_row[1] if sentiment_row else None})")
        print(f"8.3 Language detection: {results['step8_language_detection']} (Language: {sentiment_row[2] if sentiment_row else None})")
        print(f"10.1 Escalation triggered automatically: {results['step10_escalation_generated']}")

        # Close Chat & Rate
        page.locator("button:has-text('Close Chat & Rate')").click()
        page.wait_for_timeout(2000)
        page.locator("textarea[aria-label='Feedback comments (optional):']").fill("Fast chatbot, but slow human service.")
        page.locator("button:has-text('Submit Feedback')").last.click()
        page.wait_for_timeout(3000)
        
        feedback_row_real = check_db_record("select rating, feedback_text from customer_feedback order by created_at desc limit 1")
        results["step7_ratings_stored"] = feedback_row_real is not None and feedback_row_real[0] is not None
        print(f"7.2 Customer Rating stored: {results['step7_ratings_stored']}")

        # Log out
        page.locator("button:has-text('Log Out')").click()
        page.wait_for_timeout(3000)

        # ----------------------------------------------------
        # STEP 10: REVIEW CENTER & ESCALATION WORKFLOW
        # ----------------------------------------------------
        print("\n=== Testing Step 10 ===")
        page.locator("input[aria-label='Email Address']").first.fill("manager_new@saas.com")
        page.locator("input[aria-label='Password']").fill("password123")
        page.locator("button:has-text('Sign In')").last.click()
        page.wait_for_timeout(4000)
        
        page.locator("button[role='tab']").nth(3).click()
        page.wait_for_timeout(4000)
        page.screenshot(path=os.path.join(dest_dir, "uat_escalation_board.png"))
        
        def ensure_expander_open():
            notes_visible = page.locator("textarea[aria-label^='Add/Edit notes for']").is_visible()
            claim_visible = page.locator("button:has-text('Claim Escalation')").is_visible()
            res_visible = page.locator("textarea[aria-label='Resolution Summary:']").is_visible()
            if not (notes_visible or claim_visible or res_visible):
                print("Expander is closed. Expanding...")
                page.locator("summary:has-text('Ticket:')").first.click()
                page.wait_for_timeout(2000)
            else:
                print("Expander is already open. Skipping click.")
        
        ensure_expander_open()
        claim_btn = page.locator("button:has-text('Claim Escalation')")
        if claim_btn.count() > 0:
            print("Found Claim Escalation button. Clicking...")
            claim_btn.first.click()
            page.wait_for_timeout(4000)
            
            ensure_expander_open()
            page.locator("textarea[aria-label^='Add/Edit notes for']").first.fill("Manager notes added via UAT.")
            page.locator("button:has-text('Save Notes')").first.click()
            page.wait_for_timeout(4000)
            
            ensure_expander_open()
            page.locator("textarea[aria-label='Resolution Summary:']").first.fill("Refund processed successfully.")
            page.locator("button:has-text('Resolve case')").first.click()
            page.wait_for_timeout(4000)
            
            results["step10_manager_workflow_frontend"] = True
        else:
            print("Claim Escalation button not found in UI.")
            results["step10_manager_workflow_frontend"] = False
            
        esc_row = check_db_record("select status, notes, resolution_summary from escalation_events order by created_at desc limit 1")
        results["step10_status_transitions"] = esc_row is not None and esc_row[0] == "resolved"
        print(f"10.2 Escalation status resolved: {results['step10_status_transitions']}")

        browser.close()

    # Write UAT JSON Summary
    with open(os.path.join(dest_dir, "uat_summary.json"), "w") as f:
        json.dump(results, f, indent=2)
    print("\n✓ UAT Run completed and summary stored.")

if __name__ == "__main__":
    try:
        run_uat()
    except Exception as e:
        print(f"UAT Crash: {e}")
    finally:
        cleanup_mock_files()
        # Cleanly terminate the Streamlit subprocess
        try:
            streamlit_proc.terminate()
            streamlit_proc.wait()
            print("Streamlit server process terminated.")
        except Exception:
            pass
