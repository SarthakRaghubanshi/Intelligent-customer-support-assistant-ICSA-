import io
import csv
from abc import ABC, abstractmethod
from typing import BinaryIO

# Try importing parsing packages
try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None


class BaseParser(ABC):
    """
    Abstract Base Class for all document parsers.
    """
    @abstractmethod
    def parse(self, file_stream: BinaryIO) -> str:
        """
        Parses a file-like binary stream and returns clean extracted text.
        """
        pass


class PDFParser(BaseParser):
    """
    Parser for PDF documents using pypdf.
    """
    def parse(self, file_stream: BinaryIO) -> str:
        if pypdf is None:
            raise RuntimeError("pypdf library is not installed.")
        try:
            # Re-verify stream is at beginning
            file_stream.seek(0)
            reader = pypdf.PdfReader(file_stream)
            text_parts = []
            for idx, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            # Combine page contents
            full_text = "\n".join(text_parts).strip()
            return full_text
        except Exception as e:
            raise ValueError(f"Failed to parse PDF document: {str(e)}") from e


class DOCXParser(BaseParser):
    """
    Parser for Word documents (.docx) using python-docx.
    """
    def parse(self, file_stream: BinaryIO) -> str:
        if docx is None:
            raise RuntimeError("python-docx library is not installed.")
        try:
            file_stream.seek(0)
            doc = docx.Document(file_stream)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    
            # Extract table cell texts, deduplicating merged cells
            for table in doc.tables:
                for row in table.rows:
                    seen_cells = set()
                    row_texts = []
                    for cell in row.cells:
                        # cell._tc is the underlying XML element (lxml element),
                        # which is identical for merged cells in a row.
                        cell_id = getattr(cell, "_tc", cell)
                        if cell_id not in seen_cells:
                            seen_cells.add(cell_id)
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_texts.append(cell_text)
                    if row_texts:
                        text_parts.append(" | ".join(row_texts))
                        
            full_text = "\n".join(text_parts).strip()
            return full_text
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX document: {str(e)}") from e


class CSVParser(BaseParser):
    """
    Parser for CSV files using csv.DictReader.
    """
    def parse(self, file_stream: BinaryIO) -> str:
        try:
            file_stream.seek(0)
            # CSV DictReader needs a text wrapper. We attempt UTF-8 with BOM signature decode.
            content = file_stream.read()
            try:
                decoded_content = content.decode("utf-8-sig")
            except UnicodeDecodeError:
                decoded_content = content.decode("iso-8859-1")
                
            f_in = io.StringIO(decoded_content)
            reader = csv.DictReader(f_in)
            
            if not reader.fieldnames:
                raise ValueError("CSV is empty or missing headers.")
                
            text_parts = []
            for idx, row in enumerate(reader, 1):
                row_str = f"Row {idx}: " + ", ".join(
                    f"{col}={val.strip() if val else ''}" for col, val in row.items() if col
                )
                text_parts.append(row_str)
                
            full_text = "\n".join(text_parts).strip()
            return full_text
        except Exception as e:
            raise ValueError(f"Failed to parse CSV document: {str(e)}") from e


class TXTParser(BaseParser):
    """
    Parser for plain text files (.txt) with fallback encoding support.
    """
    def parse(self, file_stream: BinaryIO) -> str:
        try:
            file_stream.seek(0)
            content = file_stream.read()
            try:
                return content.decode("utf-8-sig").strip()
            except UnicodeDecodeError:
                return content.decode("iso-8859-1").strip()
        except Exception as e:
            raise ValueError(f"Failed to parse TXT document: {str(e)}") from e
